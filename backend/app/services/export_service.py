import io
import json
from typing import Any


# ── Shared helpers ────────────────────────────────────────────────────────────

def _parse_content(raw: Any) -> Any:
    if isinstance(raw, str):
        try:
            return json.loads(raw)
        except (json.JSONDecodeError, TypeError):
            return {}
    return raw or {}


def _section_map(sections: list[dict]) -> dict:
    return {s["section_type"]: _parse_content(s.get("content")) for s in sections}


def _build_pdf(story: list, margins=(0.75, 0.75, 0.75, 0.75)) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.units import inch
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=margins[0] * inch, rightMargin=margins[1] * inch,
        topMargin=margins[2] * inch, bottomMargin=margins[3] * inch,
    )
    doc.build(story)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────

def generate_pdf(resume: dict, sections: list[dict]) -> bytes:
    template = resume.get("template", "modern")
    sm = _section_map(sections)
    dispatch = {
        "classic": _pdf_classic,
        "executive": _pdf_executive,
        "compact": _pdf_compact,
        "elegant": _pdf_elegant,
    }
    return dispatch.get(template, _pdf_modern)(resume, sm)


def generate_docx(resume: dict, sections: list[dict]) -> bytes:
    template = resume.get("template", "modern")
    sm = _section_map(sections)
    dispatch = {
        "classic": _docx_classic,
        "executive": _docx_executive,
        "compact": _docx_compact,
        "elegant": _docx_elegant,
    }
    return dispatch.get(template, _docx_modern)(resume, sm)


# ═══════════════════════════════════════════════════════════════════════════════
# PDF TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════

# ── 1. Modern Minimal (default) ───────────────────────────────────────────────

def _pdf_modern(resume: dict, sm: dict) -> bytes:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER

    INDIGO = colors.HexColor("#4F46E5")
    styles = getSampleStyleSheet()
    name_s = ParagraphStyle("name", fontSize=20, leading=24, alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=3)
    contact_s = ParagraphStyle("contact", fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=6)
    head_s = ParagraphStyle("head", fontSize=10, fontName="Helvetica-Bold", textColor=INDIGO, spaceBefore=10, spaceAfter=3, leading=13)
    body_s = ParagraphStyle("body", fontSize=10, leading=13)
    sub_s = ParagraphStyle("sub", fontSize=10, fontName="Helvetica-Bold", leading=13)
    bullet_s = ParagraphStyle("bullet", fontSize=10, leading=13, leftIndent=14, firstLineIndent=-10)

    contact = sm.get("contact", {})
    story = [
        Paragraph(contact.get("full_name") or resume.get("title", "Resume"), name_s),
    ]
    parts = [p for p in [contact.get("email"), contact.get("phone"), contact.get("location"), contact.get("linkedin"), contact.get("github")] if p]
    if parts:
        story.append(Paragraph(" · ".join(parts), contact_s))
    story.append(HRFlowable(width="100%", thickness=1.5, color=INDIGO))

    _add_summary_pdf(story, sm, head_s, body_s, label="SUMMARY")
    _add_experience_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=colors.HexColor("#C7D2FE"))
    _add_education_pdf(story, sm, head_s, sub_s, body_s, hr_color=colors.HexColor("#C7D2FE"))
    _add_skills_pdf(story, sm, head_s, body_s, hr_color=colors.HexColor("#C7D2FE"))
    _add_projects_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=colors.HexColor("#C7D2FE"))

    return _build_pdf(story)


# ── 2. Classic Professional ───────────────────────────────────────────────────

def _pdf_classic(resume: dict, sm: dict) -> bytes:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER

    BLACK = colors.black
    styles = getSampleStyleSheet()
    name_s = ParagraphStyle("name", fontSize=18, leading=22, alignment=TA_CENTER, fontName="Times-Bold", spaceAfter=3)
    contact_s = ParagraphStyle("contact", fontSize=9, alignment=TA_CENTER, fontName="Times-Roman", spaceAfter=6)
    head_s = ParagraphStyle("head", fontSize=11, fontName="Times-Bold", spaceBefore=10, spaceAfter=2, leading=14)
    body_s = ParagraphStyle("body", fontSize=10, fontName="Times-Roman", leading=14)
    sub_s = ParagraphStyle("sub", fontSize=10, fontName="Times-Bold", leading=14)
    bullet_s = ParagraphStyle("bullet", fontSize=10, fontName="Times-Roman", leading=14, leftIndent=14, firstLineIndent=-10)

    contact = sm.get("contact", {})
    story = [
        Paragraph(contact.get("full_name") or resume.get("title", "Resume"), name_s),
    ]
    parts = [p for p in [contact.get("email"), contact.get("phone"), contact.get("location"), contact.get("linkedin"), contact.get("github")] if p]
    if parts:
        story.append(Paragraph(" · ".join(parts), contact_s))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLACK))

    _add_summary_pdf(story, sm, head_s, body_s, hr_color=BLACK, label="SUMMARY")
    _add_experience_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=BLACK)
    _add_education_pdf(story, sm, head_s, sub_s, body_s, hr_color=BLACK)
    _add_skills_pdf(story, sm, head_s, body_s, hr_color=BLACK)
    _add_projects_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=BLACK)

    return _build_pdf(story, margins=(1.0, 1.0, 1.0, 1.0))


# ── 3. Executive ──────────────────────────────────────────────────────────────

def _pdf_executive(resume: dict, sm: dict) -> bytes:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter

    NAVY = colors.HexColor("#1C2B3A")
    LIGHT_GRAY = colors.HexColor("#CBD5E0")
    WHITE = colors.white

    name_s = ParagraphStyle("name", fontSize=22, leading=26, alignment=TA_CENTER, fontName="Helvetica-Bold", textColor=WHITE)
    job_s = ParagraphStyle("job", fontSize=10, alignment=TA_CENTER, fontName="Helvetica", textColor=colors.HexColor("#A0AEC0"), spaceAfter=4)
    contact_s = ParagraphStyle("contact_exec", fontSize=8, alignment=TA_CENTER, fontName="Helvetica", textColor=colors.HexColor("#CBD5E0"))
    head_s = ParagraphStyle("head", fontSize=11, fontName="Helvetica-Bold", textColor=NAVY, spaceBefore=12, spaceAfter=3, leading=14)
    body_s = ParagraphStyle("body", fontSize=10, leading=13)
    sub_s = ParagraphStyle("sub", fontSize=10, fontName="Helvetica-Bold", leading=13)
    bullet_s = ParagraphStyle("bullet", fontSize=10, leading=13, leftIndent=14, firstLineIndent=-10)

    contact = sm.get("contact", {})
    name_text = contact.get("full_name") or resume.get("title", "Resume")
    target = resume.get("target_job_title", "")
    contact_parts = [p for p in [contact.get("email"), contact.get("phone"), contact.get("location")] if p]

    # Header band as a table cell
    usable_width = letter[0] - 1.5 * inch  # 0.75 margins each side
    header_content = [Paragraph(name_text, name_s)]
    if target:
        header_content.append(Paragraph(target, job_s))
    if contact_parts:
        header_content.append(Paragraph("  ·  ".join(contact_parts), contact_s))

    header_table = Table([[header_content]], colWidths=[usable_width])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), NAVY),
        ("TOPPADDING", (0, 0), (-1, -1), 18),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 18),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))

    story = [header_table, Spacer(1, 10)]

    _add_summary_pdf(story, sm, head_s, body_s, hr_color=LIGHT_GRAY, label="SUMMARY")
    _add_experience_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=LIGHT_GRAY)
    _add_education_pdf(story, sm, head_s, sub_s, body_s, hr_color=LIGHT_GRAY)
    _add_skills_pdf(story, sm, head_s, body_s, hr_color=LIGHT_GRAY)
    _add_projects_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=LIGHT_GRAY)

    return _build_pdf(story)


# ── 4. Compact Technical (two-column) ────────────────────────────────────────

def _pdf_compact(resume: dict, sm: dict) -> bytes:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter

    DARK = colors.HexColor("#374151")
    SIDEBAR_BG = colors.HexColor("#F3F4F6")
    RULE = colors.HexColor("#E5E7EB")

    # Compact styles
    name_s = ParagraphStyle("name", fontSize=16, fontName="Helvetica-Bold", textColor=DARK, spaceAfter=2)
    target_s = ParagraphStyle("target", fontSize=9, textColor=colors.HexColor("#6B7280"), spaceAfter=8)
    head_s = ParagraphStyle("head", fontSize=8, fontName="Helvetica-Bold", textColor=DARK, spaceBefore=8, spaceAfter=2, leading=11)
    body_s = ParagraphStyle("body", fontSize=8.5, leading=11)
    sub_s = ParagraphStyle("sub", fontSize=8.5, fontName="Helvetica-Bold", leading=11)
    bullet_s = ParagraphStyle("bullet", fontSize=8.5, leading=11, leftIndent=10, firstLineIndent=-8)
    sidebar_head_s = ParagraphStyle("sh", fontSize=8, fontName="Helvetica-Bold", textColor=DARK, spaceBefore=8, spaceAfter=2)
    sidebar_body_s = ParagraphStyle("sb", fontSize=8, leading=11, textColor=DARK)

    contact = sm.get("contact", {})
    skills = sm.get("skills", {})
    education = sm.get("education", []) or []
    if isinstance(education, str):
        education = []

    # ── Sidebar content ──
    sidebar = []
    # Contact
    sidebar.append(Paragraph("CONTACT", sidebar_head_s))
    sidebar.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
    for label, key in [("Email", "email"), ("Phone", "phone"), ("Location", "location"), ("LinkedIn", "linkedin"), ("GitHub", "github")]:
        val = contact.get(key, "")
        if val:
            sidebar.append(Paragraph(f"<b>{label}</b>", sidebar_body_s))
            sidebar.append(Paragraph(val, sidebar_body_s))
    # Skills
    categories = skills.get("categories", []) if isinstance(skills, dict) else []
    if categories:
        sidebar.append(Spacer(1, 6))
        sidebar.append(Paragraph("SKILLS", sidebar_head_s))
        sidebar.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
        for cat in categories:
            if isinstance(cat, dict):
                sidebar.append(Paragraph(f"<b>{cat.get('name', '')}</b>", sidebar_body_s))
                sidebar.append(Paragraph(", ".join(cat.get("items", [])), sidebar_body_s))
    # Education in sidebar
    if isinstance(education, list) and education:
        sidebar.append(Spacer(1, 6))
        sidebar.append(Paragraph("EDUCATION", sidebar_head_s))
        sidebar.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
        for edu in education:
            if isinstance(edu, dict):
                sidebar.append(Paragraph(f"<b>{edu.get('school', '')}</b>", sidebar_body_s))
                sidebar.append(Paragraph(edu.get("degree", ""), sidebar_body_s))
                sidebar.append(Paragraph(edu.get("year", ""), sidebar_body_s))

    # ── Main content ──
    main = []
    main.append(Paragraph(contact.get("full_name") or resume.get("title", "Resume"), name_s))
    if resume.get("target_job_title"):
        main.append(Paragraph(resume["target_job_title"], target_s))

    summary = sm.get("summary", {})
    summary_text = summary.get("text", "") if isinstance(summary, dict) else ""
    if summary_text:
        main.append(Paragraph("SUMMARY", head_s))
        main.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
        main.append(Paragraph(summary_text, body_s))

    experience = sm.get("experience", [])
    if isinstance(experience, list) and experience:
        main.append(Paragraph("EXPERIENCE", head_s))
        main.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
        for job in experience:
            if not isinstance(job, dict):
                continue
            main.append(Spacer(1, 4))
            main.append(Paragraph(f"{job.get('title', '')} — {job.get('company', '')}", sub_s))
            if job.get("dates"):
                main.append(Paragraph(job["dates"], body_s))
            for b in job.get("bullets", []):
                main.append(Paragraph(f"– {b}", bullet_s))

    projects = sm.get("projects", [])
    if isinstance(projects, list) and projects:
        main.append(Paragraph("PROJECTS", head_s))
        main.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            main.append(Spacer(1, 4))
            proj_line = proj.get("name", "")
            if proj.get("tech"):
                proj_line += f" — {proj['tech']}"
            main.append(Paragraph(proj_line, sub_s))
            for b in proj.get("bullets", []):
                main.append(Paragraph(f"– {b}", bullet_s))

    # Two-column table
    usable_width = letter[0] - 1.5 * inch
    sidebar_w = 2.0 * inch
    main_w = usable_width - sidebar_w

    table = Table([[sidebar, main]], colWidths=[sidebar_w, main_w])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), SIDEBAR_BG),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, -1), 8),
        ("RIGHTPADDING", (0, 0), (0, -1), 8),
        ("LEFTPADDING", (1, 0), (1, -1), 12),
        ("RIGHTPADDING", (1, 0), (1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))

    return _build_pdf([table])


# ── 5. Elegant ────────────────────────────────────────────────────────────────

def _pdf_elegant(resume: dict, sm: dict) -> bytes:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER

    NAVY = colors.HexColor("#1B2A4A")
    SOFT = colors.HexColor("#C3CFE2")
    MID = colors.HexColor("#4A5568")

    name_s = ParagraphStyle("name", fontSize=22, leading=26, alignment=TA_CENTER, fontName="Helvetica", textColor=NAVY, spaceAfter=4, tracking=2)
    contact_s = ParagraphStyle("contact", fontSize=9, alignment=TA_CENTER, fontName="Helvetica", textColor=MID, spaceAfter=8)
    head_s = ParagraphStyle("head", fontSize=9, fontName="Helvetica-Bold", textColor=NAVY, spaceBefore=14, spaceAfter=4, leading=13, alignment=TA_CENTER)
    body_s = ParagraphStyle("body", fontSize=10, leading=15, textColor=colors.HexColor("#1a1a1a"))
    sub_s = ParagraphStyle("sub", fontSize=10, fontName="Helvetica-Bold", leading=15)
    bullet_s = ParagraphStyle("bullet", fontSize=10, leading=15, leftIndent=16, firstLineIndent=-10)

    def _thin_rule():
        return HRFlowable(width="100%", thickness=0.4, color=SOFT)

    def _section_head(label):
        return [
            Paragraph(label, head_s),
            HRFlowable(width="50%", thickness=0.5, color=NAVY, spaceAfter=4),
        ]

    contact = sm.get("contact", {})
    story = [
        Spacer(1, 4),
        Paragraph(contact.get("full_name") or resume.get("title", "Resume"), name_s),
        HRFlowable(width="55%", thickness=0.5, color=NAVY, spaceAfter=2),
        HRFlowable(width="55%", thickness=0.5, color=NAVY, spaceAfter=4),
    ]
    parts = [p for p in [contact.get("email"), contact.get("phone"), contact.get("location"), contact.get("linkedin"), contact.get("github")] if p]
    if parts:
        story.append(Paragraph("  ·  ".join(parts), contact_s))
    story.append(_thin_rule())

    summary = sm.get("summary", {})
    summary_text = summary.get("text", "") if isinstance(summary, dict) else ""
    if summary_text:
        story += _section_head("S U M M A R Y")
        story.append(Paragraph(summary_text, body_s))
        story.append(_thin_rule())

    experience = sm.get("experience", [])
    if isinstance(experience, list) and experience:
        story += _section_head("E X P E R I E N C E")
        for job in experience:
            if not isinstance(job, dict):
                continue
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"{job.get('title', '')} — {job.get('company', '')}", sub_s))
            if job.get("dates"):
                story.append(Paragraph(job["dates"], body_s))
            for b in job.get("bullets", []):
                story.append(Paragraph(f"›  {b}", bullet_s))
        story.append(_thin_rule())

    education = sm.get("education", [])
    if isinstance(education, list) and education:
        story += _section_head("E D U C A T I O N")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"{edu.get('degree', '')} — {edu.get('school', '')}", sub_s))
            detail = [p for p in [edu.get("year"), edu.get("gpa") and f"GPA: {edu['gpa']}"] if p]
            if detail:
                story.append(Paragraph(", ".join(detail), body_s))
        story.append(_thin_rule())

    skills = sm.get("skills", {})
    categories = skills.get("categories", []) if isinstance(skills, dict) else []
    if categories:
        story += _section_head("S K I L L S")
        for cat in categories:
            if isinstance(cat, dict):
                story.append(Paragraph(f"<b>{cat.get('name', '')}:</b>  {', '.join(cat.get('items', []))}", body_s))
        story.append(_thin_rule())

    projects = sm.get("projects", [])
    if isinstance(projects, list) and projects:
        story += _section_head("P R O J E C T S")
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            story.append(Spacer(1, 6))
            proj_line = proj.get("name", "")
            if proj.get("tech"):
                proj_line += f" — {proj['tech']}"
            story.append(Paragraph(proj_line, sub_s))
            for b in proj.get("bullets", []):
                story.append(Paragraph(f"›  {b}", bullet_s))

    return _build_pdf(story, margins=(1.0, 1.0, 0.85, 0.85))


# ── Shared PDF section builders ───────────────────────────────────────────────

def _add_summary_pdf(story, sm, head_s, body_s, hr_color=None, label="SUMMARY"):
    from reportlab.platypus import Paragraph, HRFlowable
    from reportlab.lib import colors
    summary = sm.get("summary", {})
    text = summary.get("text", "") if isinstance(summary, dict) else ""
    if not text:
        return
    story.append(Paragraph(label, head_s))
    if hr_color:
        story.append(HRFlowable(width="100%", thickness=0.5, color=hr_color))
    story.append(Paragraph(text, body_s))


def _add_experience_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=None):
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    experience = sm.get("experience", [])
    if not isinstance(experience, list) or not experience:
        return
    story.append(Paragraph("EXPERIENCE", head_s))
    if hr_color:
        story.append(HRFlowable(width="100%", thickness=0.5, color=hr_color))
    for job in experience:
        if not isinstance(job, dict):
            continue
        story.append(Spacer(1, 5))
        story.append(Paragraph(f"{job.get('title', '')} — {job.get('company', '')}", sub_s))
        if job.get("dates"):
            story.append(Paragraph(job["dates"], body_s))
        for b in job.get("bullets", []):
            story.append(Paragraph(f"• {b}", bullet_s))


def _add_education_pdf(story, sm, head_s, sub_s, body_s, hr_color=None):
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    education = sm.get("education", [])
    if not isinstance(education, list) or not education:
        return
    story.append(Paragraph("EDUCATION", head_s))
    if hr_color:
        story.append(HRFlowable(width="100%", thickness=0.5, color=hr_color))
    for edu in education:
        if not isinstance(edu, dict):
            continue
        story.append(Spacer(1, 5))
        story.append(Paragraph(f"{edu.get('degree', '')} — {edu.get('school', '')}", sub_s))
        detail = [p for p in [edu.get("year"), edu.get("gpa") and f"GPA: {edu['gpa']}"] if p]
        if detail:
            story.append(Paragraph(", ".join(detail), body_s))


def _add_skills_pdf(story, sm, head_s, body_s, hr_color=None):
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    skills = sm.get("skills", {})
    categories = skills.get("categories", []) if isinstance(skills, dict) else []
    if not categories:
        return
    story.append(Paragraph("SKILLS", head_s))
    if hr_color:
        story.append(HRFlowable(width="100%", thickness=0.5, color=hr_color))
    for cat in categories:
        if isinstance(cat, dict):
            story.append(Paragraph(f"<b>{cat.get('name', '')}:</b>  {', '.join(cat.get('items', []))}", body_s))


def _add_projects_pdf(story, sm, head_s, sub_s, body_s, bullet_s, hr_color=None):
    from reportlab.platypus import Paragraph, Spacer, HRFlowable
    projects = sm.get("projects", [])
    if not isinstance(projects, list) or not projects:
        return
    story.append(Paragraph("PROJECTS", head_s))
    if hr_color:
        story.append(HRFlowable(width="100%", thickness=0.5, color=hr_color))
    for proj in projects:
        if not isinstance(proj, dict):
            continue
        story.append(Spacer(1, 5))
        proj_line = proj.get("name", "")
        if proj.get("tech"):
            proj_line += f" — {proj['tech']}"
        story.append(Paragraph(proj_line, sub_s))
        for b in proj.get("bullets", []):
            story.append(Paragraph(f"• {b}", bullet_s))


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════

def _docx_base(resume: dict, sm: dict, accent_rgb: tuple, use_serif: bool = False) -> bytes:
    """Shared DOCX builder used by most templates."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    margins = Pt(72) if not use_serif else Pt(80)
    sec.top_margin = sec.bottom_margin = margins
    sec.left_margin = sec.right_margin = margins

    r, g, b = accent_rgb
    font = "Times New Roman" if use_serif else "Calibri"

    def heading(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = font
        run.font.color.rgb = RGBColor(r, g, b)

    contact = sm.get("contact", {})
    name = contact.get("full_name") or resume.get("title", "Resume")
    np = doc.add_paragraph(name)
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.runs[0]
    nr.bold = True
    nr.font.size = Pt(20)
    nr.font.name = font
    nr.font.color.rgb = RGBColor(r, g, b)

    parts = [p for p in [contact.get("email"), contact.get("phone"), contact.get("location"), contact.get("linkedin"), contact.get("github")] if p]
    if parts:
        cp = doc.add_paragraph(" · ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.name = font

    summary = sm.get("summary", {})
    summary_text = summary.get("text", "") if isinstance(summary, dict) else ""
    if summary_text:
        heading("SUMMARY")
        p = doc.add_paragraph(summary_text)
        p.runs[0].font.name = font

    experience = sm.get("experience", [])
    if isinstance(experience, list) and experience:
        heading("EXPERIENCE")
        for job in experience:
            if not isinstance(job, dict):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            r2 = p.add_run(f"{job.get('title', '')} — {job.get('company', '')}")
            r2.bold = True
            r2.font.name = font
            if job.get("dates"):
                dp = doc.add_paragraph(job["dates"])
                dp.runs[0].font.name = font
                dp.runs[0].font.size = Pt(9)
            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.runs[0].font.name = font

    education = sm.get("education", [])
    if isinstance(education, list) and education:
        heading("EDUCATION")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            r2 = p.add_run(f"{edu.get('degree', '')} — {edu.get('school', '')}")
            r2.bold = True
            r2.font.name = font
            detail = [x for x in [edu.get("year"), edu.get("gpa") and f"GPA: {edu['gpa']}"] if x]
            if detail:
                dp = doc.add_paragraph(", ".join(detail))
                dp.runs[0].font.name = font

    skills = sm.get("skills", {})
    categories = skills.get("categories", []) if isinstance(skills, dict) else []
    if categories:
        heading("SKILLS")
        for cat in categories:
            if isinstance(cat, dict):
                p = doc.add_paragraph()
                r2 = p.add_run(f"{cat.get('name', '')}: ")
                r2.bold = True
                r2.font.name = font
                r3 = p.add_run(", ".join(cat.get("items", [])))
                r3.font.name = font

    projects = sm.get("projects", [])
    if isinstance(projects, list) and projects:
        heading("PROJECTS")
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            proj_line = proj.get("name", "")
            if proj.get("tech"):
                proj_line += f" — {proj['tech']}"
            r2 = p.add_run(proj_line)
            r2.bold = True
            r2.font.name = font
            for bullet in proj.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.runs[0].font.name = font

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_modern(resume, sm):
    return _docx_base(resume, sm, accent_rgb=(79, 70, 229))   # indigo


def _docx_classic(resume, sm):
    return _docx_base(resume, sm, accent_rgb=(0, 0, 0), use_serif=True)


def _docx_executive(resume, sm):
    return _docx_base(resume, sm, accent_rgb=(28, 43, 58))    # dark navy


def _docx_elegant(resume, sm):
    return _docx_base(resume, sm, accent_rgb=(27, 42, 74))    # deep navy


def _docx_compact(resume, sm):
    # python-docx doesn't support true two-column with sidebar bg,
    # so compact DOCX uses a dense single-column layout.
    return _docx_base(resume, sm, accent_rgb=(55, 65, 81))    # slate
